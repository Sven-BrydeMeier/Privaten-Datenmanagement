"""
Immobilien-Verwaltung - Verwaltung von Immobilien für automatische Dokumentenzuordnung
"""
import streamlit as st
from pathlib import Path
import sys
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from database.db import init_db, get_db, get_current_user_id
from database.models import Property, Document, Folder
from utils.components import render_sidebar_cart
from utils.helpers import format_date

st.set_page_config(page_title="Immobilien", page_icon="🏘️", layout="wide")
init_db()

# Sidebar mit Aktentasche
render_sidebar_cart()


def link_documents_to_property(session, user_id: int, property_id: int,
                                street: str, house_number: str, postal_code: str, city: str) -> int:
    """
    Durchsucht alle Dokumente nach der Immobilienadresse und verknüpft passende Dokumente.

    Returns:
        Anzahl der verknüpften Dokumente
    """
    linked_count = 0

    # Suchbegriffe erstellen
    search_terms = []

    # Vollständige Adresse
    full_address = f"{street} {house_number}".strip()
    if full_address:
        search_terms.append(full_address.lower())

    # Straße ohne Hausnummer
    if street:
        search_terms.append(street.lower())

    # PLZ + Stadt
    if postal_code:
        search_terms.append(postal_code)
    if city:
        search_terms.append(city.lower())

    # PLZ Stadt Kombination
    if postal_code and city:
        search_terms.append(f"{postal_code} {city}".lower())

    # Alle Dokumente durchsuchen
    documents = session.query(Document).filter(
        Document.user_id == user_id,
        Document.property_id.is_(None),  # Nur nicht-zugeordnete
        Document.ocr_text.isnot(None)
    ).all()

    for doc in documents:
        if not doc.ocr_text:
            continue

        text_lower = doc.ocr_text.lower()
        matches = 0

        # Prüfen, ob Adressbestandteile vorkommen
        for term in search_terms:
            if term in text_lower:
                matches += 1

        # Wenn mindestens 2 Adressbestandteile gefunden werden (z.B. Straße + PLZ oder Stadt)
        if matches >= 2:
            doc.property_id = property_id
            doc.property_address = f"{street} {house_number}, {postal_code} {city}".strip()
            linked_count += 1

    if linked_count > 0:
        session.commit()

    return linked_count


def scan_all_documents_for_properties(user_id: int) -> dict:
    """
    Durchsucht alle Dokumente nach allen Immobilienadresse und verknüpft passende.

    Returns:
        Dictionary mit property_id -> Anzahl verknüpfter Dokumente
    """
    results = {}

    with get_db() as session:
        properties = session.query(Property).filter(
            Property.user_id == user_id
        ).all()

        for prop in properties:
            count = link_documents_to_property(
                session, user_id, prop.id,
                prop.street, prop.house_number, prop.postal_code, prop.city
            )
            if count > 0:
                results[prop.id] = count

    return results

st.title("🏘️ Immobilien-Verwaltung")
st.markdown("Verwalten Sie Ihre Immobilien für automatische Dokumentenzuordnung")

user_id = get_current_user_id()


def render_property_card(prop: Property, doc_count: int):
    """Rendert eine Immobilien-Karte"""
    with st.container():
        col_icon, col_info, col_stats, col_actions = st.columns([1, 4, 2, 2])

        with col_icon:
            # Icon basierend auf Typ
            icon = "🏠"
            if prop.property_type == "Gewerbe":
                icon = "🏢"
            elif prop.property_type == "Miete":
                icon = "🏘️"
            elif prop.usage == "Vermietet":
                icon = "🏡"
            st.markdown(f"### {icon}")

        with col_info:
            st.markdown(f"**{prop.name or 'Unbenannt'}**")
            st.caption(prop.full_address or "Keine Adresse")
            if prop.property_type:
                st.caption(f"Typ: {prop.property_type} | {prop.usage or '—'}")

        with col_stats:
            st.metric("Dokumente", doc_count)

        with col_actions:
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("✏️", key=f"edit_prop_{prop.id}", help="Bearbeiten"):
                    st.session_state.edit_property_id = prop.id
                    st.rerun()
            with col_b:
                if st.button("🗑️", key=f"delete_prop_{prop.id}", help="Löschen"):
                    st.session_state.delete_property_id = prop.id
                    st.rerun()

        st.divider()


# Tabs für verschiedene Ansichten
tab_list, tab_add, tab_docs = st.tabs(["📋 Übersicht", "➕ Hinzufügen", "📄 Dokumente nach Immobilie"])

with tab_list:
    st.subheader("Ihre Immobilien")

    # Button zum Durchsuchen aller Dokumente
    col_title, col_scan = st.columns([3, 1])
    with col_scan:
        if st.button("🔍 Dokumente scannen", help="Durchsucht alle Dokumente nach Immobilienadressen"):
            with st.spinner("Durchsuche Dokumente..."):
                results = scan_all_documents_for_properties(user_id)
                if results:
                    total = sum(results.values())
                    st.success(f"✅ {total} Dokumente wurden zugeordnet!")
                    st.rerun()
                else:
                    st.info("Keine neuen Zuordnungen gefunden.")

    with get_db() as session:
        properties = session.query(Property).filter(
            Property.user_id == user_id
        ).order_by(Property.name).all()

        if properties:
            for prop in properties:
                # Anzahl der zugeordneten Dokumente
                doc_count = session.query(Document).filter(
                    Document.user_id == user_id,
                    Document.property_id == prop.id
                ).count()

                render_property_card(prop, doc_count)
        else:
            st.info("Sie haben noch keine Immobilien angelegt.")
            st.markdown("""
            ### Warum Immobilien anlegen?

            Wenn Sie Immobilien in der App hinterlegen, werden Dokumente automatisch zugeordnet:

            - **Rechnungen** mit Ihrer Immobilienadresse werden erkannt
            - **Nebenkostenabrechnungen** werden dem Objekt zugeordnet
            - **Verträge** (Strom, Gas, Wasser) werden verknüpft
            - **Versicherungen** für die Immobilie werden gruppiert

            👉 Gehen Sie zum Tab **"Hinzufügen"** um Ihre erste Immobilie anzulegen.
            """)


with tab_add:
    st.subheader("Neue Immobilie hinzufügen")

    # Prüfen ob Bearbeitung aktiv
    edit_mode = False
    edit_prop = None
    if st.session_state.get('edit_property_id'):
        with get_db() as session:
            edit_prop = session.get(Property, st.session_state.edit_property_id)
            if edit_prop:
                edit_mode = True
                st.info(f"✏️ Bearbeite: {edit_prop.name or edit_prop.full_address}")

                if st.button("❌ Bearbeitung abbrechen"):
                    del st.session_state.edit_property_id
                    st.rerun()

    with st.form("property_form"):
        col1, col2 = st.columns(2)

        with col1:
            name = st.text_input(
                "Kurzname *",
                value=edit_prop.name if edit_prop else "",
                placeholder="z.B. Mietwohnung Berlin, Ferienhaus Ostsee",
                help="Ein eindeutiger Name für die Immobilie"
            )

            property_type = st.selectbox(
                "Art der Immobilie",
                options=["Eigentum", "Miete", "Gewerbe", "Ferienimmobilie", "Sonstige"],
                index=["Eigentum", "Miete", "Gewerbe", "Ferienimmobilie", "Sonstige"].index(edit_prop.property_type) if edit_prop and edit_prop.property_type else 0
            )

            usage = st.selectbox(
                "Nutzung",
                options=["Selbstgenutzt", "Vermietet", "Teilvermietet", "Leerstand"],
                index=["Selbstgenutzt", "Vermietet", "Teilvermietet", "Leerstand"].index(edit_prop.usage) if edit_prop and edit_prop.usage else 0
            )

        with col2:
            street = st.text_input(
                "Straße *",
                value=edit_prop.street if edit_prop else "",
                placeholder="Musterstraße"
            )

            house_number = st.text_input(
                "Hausnummer",
                value=edit_prop.house_number if edit_prop else "",
                placeholder="123a"
            )

        col3, col4 = st.columns(2)

        with col3:
            postal_code = st.text_input(
                "PLZ *",
                value=edit_prop.postal_code if edit_prop else "",
                placeholder="12345",
                max_chars=5
            )

        with col4:
            city = st.text_input(
                "Stadt *",
                value=edit_prop.city if edit_prop else "",
                placeholder="Berlin"
            )

        st.markdown("---")
        st.markdown("**Zusätzliche Informationen (optional)**")

        col5, col6 = st.columns(2)

        with col5:
            owner = st.text_input(
                "Eigentümer / Vermieter",
                value=edit_prop.owner if edit_prop else "",
                placeholder="Name des Eigentümers oder Vermieters"
            )

            management = st.text_input(
                "Hausverwaltung",
                value=edit_prop.management if edit_prop else "",
                placeholder="Name der Hausverwaltung"
            )

        with col6:
            acquired_date = st.date_input(
                "Einzugs-/Kaufdatum",
                value=edit_prop.acquired_date.date() if edit_prop and edit_prop.acquired_date else None
            )

            sold_date = st.date_input(
                "Auszugs-/Verkaufsdatum",
                value=edit_prop.sold_date.date() if edit_prop and edit_prop.sold_date else None
            )

        notes = st.text_area(
            "Notizen",
            value=edit_prop.notes if edit_prop else "",
            placeholder="Zusätzliche Informationen zur Immobilie...",
            height=100
        )

        submit_label = "💾 Änderungen speichern" if edit_mode else "➕ Immobilie hinzufügen"
        submitted = st.form_submit_button(submit_label, type="primary", use_container_width=True)

        if submitted:
            if not name or not street or not postal_code or not city:
                st.error("Bitte füllen Sie alle Pflichtfelder (*) aus.")
            else:
                with get_db() as session:
                    if edit_mode:
                        # Bestehende Immobilie aktualisieren
                        prop = session.get(Property, st.session_state.edit_property_id)
                        if prop:
                            prop.name = name
                            prop.street = street
                            prop.house_number = house_number
                            prop.postal_code = postal_code
                            prop.city = city
                            prop.property_type = property_type
                            prop.usage = usage
                            prop.owner = owner
                            prop.management = management
                            prop.acquired_date = datetime.combine(acquired_date, datetime.min.time()) if acquired_date else None
                            prop.sold_date = datetime.combine(sold_date, datetime.min.time()) if sold_date else None
                            prop.notes = notes
                            session.commit()
                            st.success(f"✅ Immobilie '{name}' wurde aktualisiert!")
                            del st.session_state.edit_property_id
                            st.rerun()
                    else:
                        # Neue Immobilie erstellen
                        new_prop = Property(
                            user_id=user_id,
                            name=name,
                            street=street,
                            house_number=house_number,
                            postal_code=postal_code,
                            city=city,
                            property_type=property_type,
                            usage=usage,
                            owner=owner,
                            management=management,
                            acquired_date=datetime.combine(acquired_date, datetime.min.time()) if acquired_date else None,
                            sold_date=datetime.combine(sold_date, datetime.min.time()) if sold_date else None,
                            notes=notes
                        )
                        session.add(new_prop)
                        session.commit()
                        st.success(f"✅ Immobilie '{name}' wurde hinzugefügt!")

                        # Ordner für Immobilie erstellen
                        immobilien_folder = session.query(Folder).filter(
                            Folder.user_id == user_id,
                            Folder.name == "Immobilien"
                        ).first()

                        if not immobilien_folder:
                            immobilien_folder = Folder(
                                user_id=user_id,
                                name="Immobilien",
                                color="#795548",
                                icon="🏘️"
                            )
                            session.add(immobilien_folder)
                            session.flush()

                        # Unterordner für diese Immobilie
                        prop_folder = Folder(
                            user_id=user_id,
                            name=name,
                            parent_id=immobilien_folder.id,
                            color="#A1887F"
                        )
                        session.add(prop_folder)
                        session.flush()

                        # Virtuelle Unterordner für Immobilienkategorien erstellen
                        property_subfolders = [
                            ("📜 Grundsteuer", "#8D6E63"),
                            ("🛡️ Versicherung", "#7986CB"),
                            ("💰 Hausgeldabrechnung", "#4DB6AC"),
                            ("📊 Wirtschaftsplan", "#FFB74D"),
                            ("🔧 Sanierungen & Reparaturen", "#FF8A65"),
                            ("📋 Verträge", "#9575CD"),
                            ("📄 Korrespondenz", "#90A4AE"),
                        ]

                        for subfolder_name, color in property_subfolders:
                            subfolder = Folder(
                                user_id=user_id,
                                name=subfolder_name,
                                parent_id=prop_folder.id,
                                color=color
                            )
                            session.add(subfolder)

                        session.commit()
                        st.info(f"📁 Ordner 'Immobilien/{name}' mit Unterordnern wurde erstellt.")

                        # Dokumente nach Immobilienadresse durchsuchen
                        st.info("🔍 Durchsuche vorhandene Dokumente nach Adresse...")
                        property_id = new_prop.id
                        linked_count = link_documents_to_property(
                            session, user_id, property_id,
                            street, house_number, postal_code, city
                        )

                        if linked_count > 0:
                            st.success(f"✅ {linked_count} Dokumente wurden automatisch mit der Immobilie verknüpft!")
                        else:
                            st.info("Keine passenden Dokumente gefunden.")

                        st.rerun()


with tab_docs:
    st.subheader("Dokumente nach Immobilie")

    # Prüfen ob Dokumentenansicht aktiv ist
    if st.session_state.get('view_property_doc_id'):
        view_doc_id = st.session_state.view_property_doc_id

        with get_db() as session:
            doc = session.get(Document, view_doc_id)
            if doc:
                # Zurück-Button
                if st.button("⬅️ Zurück zur Liste"):
                    del st.session_state.view_property_doc_id
                    st.rerun()

                st.divider()
                st.subheader(f"📄 {doc.title or doc.filename}")

                # Zusammenfassung
                if doc.ai_summary:
                    st.info(f"📝 **Zusammenfassung:** {doc.ai_summary}")

                # Dokument-Info und Aktionen
                col_info, col_actions = st.columns([2, 1])

                with col_info:
                    st.markdown("### 📋 Dokument-Details")
                    st.write(f"**Absender:** {doc.sender or '—'}")
                    st.write(f"**Kategorie:** {doc.category or '—'}")
                    st.write(f"**Datum:** {format_date(doc.document_date)}")
                    if doc.invoice_amount:
                        st.write(f"**Betrag:** {doc.invoice_amount:.2f} €")
                    if doc.iban:
                        st.code(doc.iban, language=None)

                    # OCR-Text anzeigen
                    if doc.ocr_text:
                        with st.expander("📝 Erkannter Text", expanded=False):
                            st.text_area("OCR-Text", doc.ocr_text, height=300, disabled=True)

                    # Dokument-Vorschau (PDF, Excel, Bilder)
                    st.markdown("### 👁️ Dokument-Vorschau")
                    from utils.helpers import get_document_file_content, document_file_exists
                    from services.encryption import get_encryption_service
                    import base64

                    if doc.file_path and document_file_exists(doc.file_path):
                        try:
                            success, result = get_document_file_content(doc.file_path, user_id)
                            if success:
                                # Entschlüsseln nur wenn verschlüsselt
                                if doc.is_encrypted and doc.encryption_iv:
                                    encryption = get_encryption_service()
                                    try:
                                        file_data = encryption.decrypt_file(result, doc.encryption_iv, doc.filename)
                                    except:
                                        file_data = result
                                else:
                                    file_data = result

                                mime_type = doc.mime_type or ""
                                filename_lower = doc.filename.lower() if doc.filename else ""

                                # PDF-Vorschau
                                if mime_type == "application/pdf" or filename_lower.endswith(".pdf"):
                                    pdf_base64 = base64.b64encode(file_data).decode('utf-8')
                                    pdf_display = f'''
                                    <iframe
                                        src="data:application/pdf;base64,{pdf_base64}"
                                        width="100%"
                                        height="600px"
                                        type="application/pdf"
                                        style="border: 1px solid #ddd; border-radius: 5px;">
                                    </iframe>
                                    '''
                                    st.markdown(pdf_display, unsafe_allow_html=True)

                                # Excel-Vorschau
                                elif filename_lower.endswith((".xlsx", ".xls")) or "spreadsheet" in mime_type:
                                    try:
                                        import pandas as pd
                                        import io

                                        # Excel lesen
                                        excel_file = io.BytesIO(file_data)

                                        # Alle Sheets auflisten
                                        xl = pd.ExcelFile(excel_file)
                                        sheet_names = xl.sheet_names

                                        if len(sheet_names) > 1:
                                            selected_sheet = st.selectbox(
                                                "Tabellenblatt auswählen",
                                                sheet_names,
                                                key=f"sheet_select_{doc.id}"
                                            )
                                        else:
                                            selected_sheet = sheet_names[0]

                                        # Sheet laden und anzeigen
                                        df = pd.read_excel(excel_file, sheet_name=selected_sheet)
                                        st.dataframe(df, use_container_width=True, height=400)

                                        st.caption(f"📊 {len(df)} Zeilen × {len(df.columns)} Spalten")
                                    except Exception as excel_err:
                                        st.warning(f"Excel-Vorschau nicht möglich: {excel_err}")

                                # Word-Vorschau
                                elif filename_lower.endswith(".docx"):
                                    try:
                                        from docx import Document as DocxDocument
                                        import io

                                        docx_file = io.BytesIO(file_data)
                                        doc_content = DocxDocument(docx_file)

                                        # Absätze extrahieren
                                        full_text = []
                                        for para in doc_content.paragraphs:
                                            if para.text.strip():
                                                # Überschriften hervorheben
                                                if para.style and para.style.name.startswith('Heading'):
                                                    full_text.append(f"\n### {para.text}\n")
                                                else:
                                                    full_text.append(para.text)

                                        # Tabellen extrahieren
                                        if doc_content.tables:
                                            full_text.append("\n---\n**Tabellen:**\n")
                                            for table in doc_content.tables:
                                                table_data = []
                                                for row in table.rows:
                                                    row_data = [cell.text.strip() for cell in row.cells]
                                                    table_data.append(" | ".join(row_data))
                                                full_text.append("\n".join(table_data))
                                                full_text.append("\n")

                                        text_content = "\n".join(full_text)
                                        st.markdown(text_content)
                                    except ImportError:
                                        st.warning("python-docx nicht installiert. Bitte installieren: pip install python-docx")
                                    except Exception as word_err:
                                        st.warning(f"Word-Vorschau nicht möglich: {word_err}")
                                        if doc.ocr_text:
                                            st.info("OCR-Text wird als Fallback angezeigt")
                                            st.text_area("OCR-Text", doc.ocr_text, height=300, disabled=True)

                                # Ältere .doc Dateien
                                elif filename_lower.endswith(".doc"):
                                    st.info("📄 Älteres Word-Format (.doc) - Bitte herunterladen und in Word öffnen")
                                    if doc.ocr_text:
                                        with st.expander("OCR-Text anzeigen"):
                                            st.text_area("OCR-Text", doc.ocr_text, height=300, disabled=True)

                                # Bild-Vorschau
                                elif mime_type.startswith("image/") or filename_lower.endswith((".jpg", ".jpeg", ".png", ".gif", ".webp")):
                                    from PIL import Image
                                    import io
                                    img = Image.open(io.BytesIO(file_data))
                                    st.image(img, use_container_width=True)

                                # Textdateien
                                elif mime_type.startswith("text/") or filename_lower.endswith((".txt", ".csv", ".json", ".xml")):
                                    try:
                                        text_content = file_data.decode('utf-8')
                                        st.code(text_content, language=None)
                                    except:
                                        st.warning("Textdatei konnte nicht dekodiert werden")

                                else:
                                    st.info(f"📄 Vorschau für {mime_type or 'unbekanntes Format'} nicht verfügbar. Bitte herunterladen.")

                        except Exception as e:
                            st.warning(f"Vorschau nicht verfügbar: {e}")
                    else:
                        st.warning("Datei nicht gefunden")

                with col_actions:
                    st.markdown("### ⚡ Aktionen")

                    # Download
                    from utils.helpers import get_document_file_content, document_file_exists
                    from services.encryption import get_encryption_service

                    if doc.file_path and document_file_exists(doc.file_path):
                        try:
                            success, result = get_document_file_content(doc.file_path, user_id)
                            if success:
                                # Entschlüsseln nur wenn verschlüsselt
                                if doc.is_encrypted and doc.encryption_iv:
                                    encryption = get_encryption_service()
                                    try:
                                        file_data = encryption.decrypt_file(result, doc.encryption_iv, doc.filename)
                                    except:
                                        file_data = result
                                else:
                                    file_data = result

                                st.download_button(
                                    "⬇️ Herunterladen",
                                    data=file_data,
                                    file_name=doc.filename,
                                    mime=doc.mime_type or "application/octet-stream",
                                    use_container_width=True
                                )
                        except Exception as e:
                            st.warning(f"Datei nicht verfügbar: {e}")
                    else:
                        st.warning("Datei nicht gefunden")

                    # Aktentasche
                    if st.button("📋 In Aktentasche", use_container_width=True):
                        if 'active_cart_items' not in st.session_state:
                            st.session_state.active_cart_items = []
                        if view_doc_id not in st.session_state.active_cart_items:
                            st.session_state.active_cart_items.append(view_doc_id)
                            st.success("✅ Zur Aktentasche hinzugefügt!")
                        else:
                            st.info("Bereits in der Aktentasche")

                    # Teilen
                    st.markdown("---")
                    st.markdown("**📤 Teilen**")

                    share_title = doc.title or doc.filename
                    share_lines = [f"📄 {share_title}"]
                    if doc.sender:
                        share_lines.append(f"Von: {doc.sender}")
                    if doc.category:
                        share_lines.append(f"Kategorie: {doc.category}")
                    if doc.document_date:
                        share_lines.append(f"Datum: {format_date(doc.document_date)}")
                    if doc.invoice_amount:
                        share_lines.append(f"Betrag: {doc.invoice_amount:.2f} €")
                    share_text = "\n".join(share_lines)

                    from utils.helpers import render_share_buttons
                    render_share_buttons(share_title, share_text, key_prefix=f"prop_doc_{view_doc_id}")

                    # Drucken
                    st.markdown("---")
                    if st.button("🖨️ Drucken", use_container_width=True):
                        st.markdown("""
                        <script>
                        window.print();
                        </script>
                        """, unsafe_allow_html=True)
                        st.info("Druckdialog wird geöffnet...")

                    # Zur Dokumentenseite wechseln
                    st.markdown("---")
                    if st.button("📄 Im Dokumentenbereich öffnen", use_container_width=True):
                        st.session_state.view_document_id = view_doc_id
                        st.switch_page("pages/3_📁_Dokumente.py")
            else:
                st.error("Dokument nicht gefunden")
                del st.session_state.view_property_doc_id
                st.rerun()
    else:
        # Normale Listenansicht
        with get_db() as session:
            properties = session.query(Property).filter(
                Property.user_id == user_id
            ).order_by(Property.name).all()

            if properties:
                # Immobilie auswählen
                prop_options = {0: "-- Alle Immobilien --"}
                prop_options.update({p.id: f"🏠 {p.name or p.full_address}" for p in properties})

                selected_prop_id = st.selectbox(
                    "Immobilie auswählen",
                    options=list(prop_options.keys()),
                    format_func=lambda x: prop_options.get(x, "")
                )

                # Dokumente laden
                query = session.query(Document).filter(
                    Document.user_id == user_id,
                    Document.property_id.isnot(None)
                )

                if selected_prop_id:
                    query = query.filter(Document.property_id == selected_prop_id)

                documents = query.order_by(Document.document_date.desc()).limit(100).all()

                if documents:
                    st.write(f"**{len(documents)} Dokumente** gefunden")

                    for doc in documents:
                        col_doc, col_prop, col_date, col_actions = st.columns([3, 2, 1, 2])

                        with col_doc:
                            st.write(f"📄 {doc.title or doc.filename}")
                            if doc.sender:
                                st.caption(f"Von: {doc.sender}")

                        with col_prop:
                            if doc.property:
                                st.caption(f"🏠 {doc.property.name or doc.property.city}")
                            if doc.property_address:
                                st.caption(f"📍 {doc.property_address[:40]}...")

                        with col_date:
                            st.caption(format_date(doc.document_date))
                            if doc.category:
                                st.caption(doc.category)

                        with col_actions:
                            btn_col1, btn_col2, btn_col3 = st.columns(3)
                            with btn_col1:
                                if st.button("👁️", key=f"view_pdoc_{doc.id}", help="Anzeigen"):
                                    st.session_state.view_property_doc_id = doc.id
                                    st.rerun()
                            with btn_col2:
                                if st.button("📋", key=f"cart_pdoc_{doc.id}", help="In Aktentasche"):
                                    if 'active_cart_items' not in st.session_state:
                                        st.session_state.active_cart_items = []
                                    if doc.id not in st.session_state.active_cart_items:
                                        st.session_state.active_cart_items.append(doc.id)
                                        st.toast("✅ Zur Aktentasche hinzugefügt!")
                                    else:
                                        st.toast("Bereits in der Aktentasche")
                            with btn_col3:
                                if st.button("📄", key=f"goto_pdoc_{doc.id}", help="Im Dokumentenbereich öffnen"):
                                    st.session_state.view_document_id = doc.id
                                    st.switch_page("pages/3_📁_Dokumente.py")

                        st.divider()
                else:
                    st.info("Keine Dokumente mit Immobilien-Zuordnung gefunden.")
            else:
                st.info("Legen Sie zuerst Immobilien an, um Dokumente zuordnen zu können.")


# Lösch-Dialog
if st.session_state.get('delete_property_id'):
    with get_db() as session:
        prop = session.get(Property, st.session_state.delete_property_id)
        if prop:
            st.warning(f"⚠️ Möchten Sie die Immobilie **'{prop.name}'** wirklich löschen?")
            st.caption("Die zugeordneten Dokumente bleiben erhalten, verlieren aber ihre Immobilien-Zuordnung.")

            col_confirm, col_cancel = st.columns(2)

            with col_confirm:
                if st.button("🗑️ Ja, löschen", type="primary", use_container_width=True):
                    # Dokumente von Immobilie trennen
                    session.query(Document).filter(
                        Document.property_id == prop.id
                    ).update({Document.property_id: None})

                    # Immobilie löschen
                    session.delete(prop)
                    session.commit()
                    st.success("Immobilie wurde gelöscht.")
                    del st.session_state.delete_property_id
                    st.rerun()

            with col_cancel:
                if st.button("❌ Abbrechen", use_container_width=True):
                    del st.session_state.delete_property_id
                    st.rerun()
