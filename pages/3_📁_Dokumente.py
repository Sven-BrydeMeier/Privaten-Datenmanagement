"""
Dokumentenmanagement - Ordnerstruktur und Dokumentenverwaltung
"""
import streamlit as st
import io
from pathlib import Path
import sys
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent))

from database.db import init_db, get_db, get_current_user_id
from database.models import Document, Folder, DocumentStatus, InvoiceStatus
from config.settings import DOCUMENT_CATEGORIES
from services.encryption import get_encryption_service
from services.document_classifier import get_classifier
from services.search_service import get_search_service
from utils.helpers import format_currency, format_date, generate_share_link, truncate_text
from utils.components import render_sidebar_cart, add_to_cart

st.set_page_config(page_title="Dokumente", page_icon="ğŸ“", layout="wide")
init_db()


def build_folder_tree(session, user_id: int, include_root: bool = False) -> list:
    """
    Baut eine hierarchische Ordnerliste fÃ¼r Selectboxen.

    Returns:
        Liste von Dicts mit 'id', 'display_name', 'path', 'depth'
    """
    from database.models import Folder

    result = []

    # Alle Ordner laden
    all_folders = session.query(Folder).filter(
        Folder.user_id == user_id
    ).order_by(Folder.name).all()

    # Index nach ID und parent_id
    folders_by_id = {f.id: f for f in all_folders}
    children_by_parent = {}
    root_folders = []

    for folder in all_folders:
        if folder.parent_id is None:
            root_folders.append(folder)
        else:
            if folder.parent_id not in children_by_parent:
                children_by_parent[folder.parent_id] = []
            children_by_parent[folder.parent_id].append(folder)

    # Sortiere Root-Ordner: Posteingang zuerst, dann alphabetisch, Papierkorb zuletzt
    def sort_key(f):
        if f.name == "Posteingang":
            return (0, f.name)
        elif f.name == "Papierkorb":
            return (2, f.name)
        else:
            return (1, f.name)

    root_folders.sort(key=sort_key)

    def add_folder_recursive(folder, depth=0, path_parts=None):
        if path_parts is None:
            path_parts = []

        current_path = path_parts + [folder.name]

        # Icon basierend auf Ordnername
        if folder.name == "Posteingang":
            icon = "ğŸ“¥"
        elif folder.name == "Papierkorb":
            icon = "ğŸ—‘ï¸"
        elif folder.name == "Archiv" or "Archiv" in folder.name:
            icon = "ğŸ“¦"
        else:
            icon = "ğŸ“‚"

        # EinrÃ¼ckung mit Baumstruktur
        if depth == 0:
            prefix = ""
        else:
            prefix = "    " * (depth - 1) + "â””â”€â”€ "

        display_name = f"{prefix}{icon} {folder.name}"
        full_path = " / ".join(current_path)

        result.append({
            'id': folder.id,
            'name': folder.name,
            'display_name': display_name,
            'path': full_path,
            'depth': depth
        })

        # Unterordner rekursiv hinzufÃ¼gen
        if folder.id in children_by_parent:
            children = sorted(children_by_parent[folder.id], key=lambda f: f.name)
            for child in children:
                add_folder_recursive(child, depth + 1, current_path)

    # Wurzelordner durchgehen
    for folder in root_folders:
        add_folder_recursive(folder)

    return result


# Sidebar mit Aktentasche
render_sidebar_cart()

user_id = get_current_user_id()

st.title("ğŸ“ Dokumente & Ordner")

# Layout: Sidebar fÃ¼r Ordner, Hauptbereich fÃ¼r Dokumente
col_folders, col_docs = st.columns([1, 3])

with col_folders:
    st.subheader("ğŸ“‚ Ordner")

    # Aktuellen Ordner aus Session
    current_folder_id = st.session_state.get('current_folder_id')

    # Ordnerdaten laden (als einfache Dicts, um DetachedInstanceError zu vermeiden)
    folder_data = []
    with get_db() as session:
        # Alle Ordner laden
        folders = session.query(Folder).filter(
            Folder.user_id == user_id,
            Folder.parent_id.is_(None)  # Nur Root-Ordner
        ).order_by(Folder.name).all()

        # Ordnerdaten extrahieren wÃ¤hrend Session aktiv ist
        for folder in folders:
            doc_count = session.query(Document).filter(
                Document.folder_id == folder.id
            ).count()

            # Unterordner laden
            subfolders_data = []
            subfolders = session.query(Folder).filter(
                Folder.parent_id == folder.id
            ).all()
            for sub in subfolders:
                sub_count = session.query(Document).filter(
                    Document.folder_id == sub.id
                ).count()
                subfolders_data.append({
                    'id': sub.id,
                    'name': sub.name,
                    'count': sub_count
                })

            folder_data.append({
                'id': folder.id,
                'name': folder.name,
                'count': doc_count,
                'subfolders': subfolders_data
            })

    # "Alle Dokumente" Option
    if st.button("ğŸ“„ Alle Dokumente", use_container_width=True,
                 type="primary" if current_folder_id is None else "secondary"):
        st.session_state.current_folder_id = None
        st.rerun()

    st.divider()

    # Ordner anzeigen - NUR ORDNER MIT DOKUMENTEN (auÃŸer Systemordner)
    for folder in folder_data:
        # Systemordner immer anzeigen, andere nur wenn sie Dokumente haben
        is_system = folder['name'] in ["Posteingang", "Papierkorb", "Archiv"]
        total_count = folder['count'] + sum(sub['count'] for sub in folder['subfolders'])

        # Ordner Ã¼berspringen wenn leer und kein Systemordner
        if not is_system and total_count == 0:
            continue

        icon = "ğŸ“¥" if folder['name'] == "Posteingang" else "ğŸ“‚"
        if folder['name'] == "Papierkorb":
            icon = "ğŸ—‘ï¸"
        elif folder['name'] == "Archiv":
            icon = "ğŸ“¦"

        is_selected = current_folder_id == folder['id']
        if st.button(f"{icon} {folder['name']} ({folder['count']})",
                    use_container_width=True,
                    type="primary" if is_selected else "secondary",
                    key=f"folder_{folder['id']}"):
            st.session_state.current_folder_id = folder['id']
            st.rerun()

        # Unterordner - nur anzeigen wenn sie Dokumente haben
        for sub in folder['subfolders']:
            if sub['count'] == 0:
                continue  # Leere Unterordner Ã¼berspringen
            if st.button(f"  â”” {sub['name']} ({sub['count']})",
                        use_container_width=True,
                        key=f"folder_{sub['id']}"):
                st.session_state.current_folder_id = sub['id']
                st.rerun()

    st.divider()

    # Neuen Ordner erstellen
    with st.expander("â• Neuer Ordner"):
        new_folder_name = st.text_input("Ordnername", key="new_folder_name")

        # Hierarchische Ordnerauswahl fÃ¼r Parent
        with get_db() as session:
            parent_tree = build_folder_tree(session, user_id)

        parent_options = [None] + [f['id'] for f in parent_tree]
        parent_folder = st.selectbox(
            "Ãœbergeordneter Ordner",
            options=parent_options,
            format_func=lambda x: "ğŸ“ Kein (Root-Ordner)" if x is None else next((f['display_name'] for f in parent_tree if f['id'] == x), ""),
            key="parent_folder"
        )

        # Pfad anzeigen
        if parent_folder:
            selected_parent = next((f for f in parent_tree if f['id'] == parent_folder), None)
            if selected_parent:
                st.caption(f"ğŸ“ Wird erstellt unter: {selected_parent['path']}")

        if st.button("Erstellen") and new_folder_name:
            with get_db() as session:
                new_folder = Folder(
                    user_id=user_id,
                    name=new_folder_name,
                    parent_id=parent_folder
                )
                session.add(new_folder)
                session.commit()
            st.success(f"âœ“ Ordner '{new_folder_name}' erstellt!")
            st.rerun()


with col_docs:
    # Breadcrumb-Navigation
    if current_folder_id:
        with get_db() as session:
            # Pfad zum aktuellen Ordner aufbauen
            breadcrumb_parts = []
            folder = session.get(Folder, current_folder_id)
            while folder:
                breadcrumb_parts.insert(0, {'id': folder.id, 'name': folder.name})
                folder = session.get(Folder, folder.parent_id) if folder.parent_id else None

            # Breadcrumb anzeigen
            bc_cols = st.columns([1] + [1] * len(breadcrumb_parts) + [4])

            with bc_cols[0]:
                if st.button("ğŸ ", key="bc_home", help="Alle Dokumente"):
                    st.session_state.current_folder_id = None
                    st.rerun()

            for i, part in enumerate(breadcrumb_parts):
                with bc_cols[i + 1]:
                    is_current = (i == len(breadcrumb_parts) - 1)
                    if is_current:
                        st.markdown(f"**ğŸ“‚ {part['name']}**")
                    else:
                        if st.button(f"ğŸ“ {part['name']}", key=f"bc_{part['id']}"):
                            st.session_state.current_folder_id = part['id']
                            st.rerun()

            st.markdown("---")

    # Suchleiste
    search_col, filter_col = st.columns([3, 1])

    with search_col:
        search_query = st.text_input("ğŸ” Suchen...", placeholder="Stichwort, Betrag, IBAN...")

    with filter_col:
        filter_category = st.selectbox(
            "Kategorie",
            options=["Alle"] + DOCUMENT_CATEGORIES,
            key="filter_category"
        )

    # Dokumente laden
    with get_db() as session:
        query = session.query(Document).filter(Document.user_id == user_id)

        # GelÃ¶schte Dokumente ausschlieÃŸen (auÃŸer im Papierkorb-Modus)
        is_trash_view = False
        if current_folder_id:
            folder = session.get(Folder, current_folder_id)
            if folder and folder.name == "Papierkorb":
                is_trash_view = True
                # Im Papierkorb: nur gelÃ¶schte Dokumente zeigen
                query = query.filter(Document.is_deleted == True)
            else:
                query = query.filter(Document.folder_id == current_folder_id)
                query = query.filter((Document.is_deleted == False) | (Document.is_deleted == None))
        else:
            # Alle Dokumente: keine gelÃ¶schten
            query = query.filter((Document.is_deleted == False) | (Document.is_deleted == None))

        # Kategoriefilter
        if filter_category != "Alle":
            query = query.filter(Document.category == filter_category)

        # Suche
        if search_query:
            search_service = get_search_service(user_id)
            search_results = search_service.search(search_query)
            doc_ids = [item['id'] for item in search_results['items']]
            if doc_ids:
                query = query.filter(Document.id.in_(doc_ids))
            else:
                # Fallback: einfache Textsuche
                query = query.filter(
                    Document.ocr_text.ilike(f'%{search_query}%') |
                    Document.filename.ilike(f'%{search_query}%') |
                    Document.sender.ilike(f'%{search_query}%')
                )

        documents = query.order_by(Document.created_at.desc()).limit(50).all()

        # Aktuellen Ordnernamen anzeigen
        if current_folder_id:
            folder = session.get(Folder, current_folder_id)
            st.subheader(f"ğŸ“‚ {folder.name}" if folder else "Dokumente")
        else:
            st.subheader("ğŸ“„ Alle Dokumente")

        st.caption(f"{len(documents)} Dokumente")

        # Dokumentenliste
        if documents:
            for doc in documents:
                with st.container():
                    col1, col2, col3, col4 = st.columns([3, 1, 1, 1])

                    with col1:
                        # Status-Icon
                        if doc.status == DocumentStatus.COMPLETED:
                            status = "âœ“"
                        elif doc.status == DocumentStatus.PROCESSING:
                            status = "â³"
                        elif doc.status == DocumentStatus.ERROR:
                            status = "âŒ"
                        else:
                            status = "ğŸ“„"

                        st.markdown(f"**{status} {doc.title or doc.filename}**")
                        meta_parts = []
                        if doc.sender:
                            meta_parts.append(doc.sender)
                        if doc.category:
                            meta_parts.append(doc.category)
                        if doc.document_date:
                            meta_parts.append(format_date(doc.document_date))
                        st.caption(" | ".join(meta_parts) if meta_parts else "Keine Metadaten")

                    with col2:
                        if doc.invoice_amount:
                            st.markdown(f"**{format_currency(doc.invoice_amount)}**")
                            if doc.invoice_status == InvoiceStatus.OPEN:
                                st.caption("ğŸ”´ Offen")
                            elif doc.invoice_status == InvoiceStatus.PAID:
                                st.caption("âœ… Bezahlt")

                    with col3:
                        if doc.iban:
                            st.code(doc.iban[:12] + "...")

                    with col4:
                        # AktionsmenÃ¼
                        with st.popover("â‹®"):
                            if st.button("ğŸ‘ï¸ Anzeigen", key=f"view_{doc.id}"):
                                st.session_state.view_document_id = doc.id
                                st.rerun()

                            if st.button("ğŸ“‹ In Aktentasche", key=f"cart_{doc.id}"):
                                if 'active_cart_items' not in st.session_state:
                                    st.session_state.active_cart_items = []
                                if doc.id not in st.session_state.active_cart_items:
                                    st.session_state.active_cart_items.append(doc.id)
                                    st.toast("âœ… Zur Aktentasche hinzugefÃ¼gt!")
                                    st.rerun()

                            if st.button("ğŸ“‚ Verschieben", key=f"move_{doc.id}"):
                                st.session_state.move_document_id = doc.id
                                st.rerun()

                            if st.button("ğŸ”— Teilen", key=f"share_{doc.id}"):
                                link = generate_share_link(doc.id)
                                st.code(link)

                            if st.button("ğŸ—‘ï¸ LÃ¶schen", key=f"del_{doc.id}"):
                                st.session_state.delete_document_id = doc.id
                                st.rerun()

                    st.divider()
        else:
            st.info("Keine Dokumente gefunden")

# Dokument anzeigen Dialog - Erweitert
if 'view_document_id' in st.session_state:
    doc_id = st.session_state.view_document_id

    with get_db() as session:
        doc = session.get(Document, doc_id)
        if doc:
            # Dokumentdaten in Dict extrahieren (fÃ¼r Verwendung auÃŸerhalb der Session)
            doc_data = {
                'id': doc.id,
                'title': doc.title,
                'filename': doc.filename,
                'file_path': doc.file_path,
                'mime_type': doc.mime_type,
                'is_encrypted': doc.is_encrypted,
                'encryption_iv': doc.encryption_iv,
                'category': doc.category,
                'sender': doc.sender,
                'sender_address': doc.sender_address,
                'document_date': doc.document_date,
                'subject': doc.subject,
                'ai_summary': doc.ai_summary,
                'reference_number': doc.reference_number,
                'customer_number': doc.customer_number,
                'insurance_number': doc.insurance_number,
                'processing_number': doc.processing_number,
                'contract_number': doc.contract_number,
                'invoice_number': getattr(doc, 'invoice_number', None),
                'invoice_amount': doc.invoice_amount,
                'invoice_due_date': doc.invoice_due_date,
                'invoice_status': doc.invoice_status,
                'invoice_paid_date': doc.invoice_paid_date,
                'paid_with_bank_account': getattr(doc, 'paid_with_bank_account', None),
                'iban': doc.iban,
                'bic': doc.bic,
                'bank_name': getattr(doc, 'bank_name', None),
                'ocr_text': doc.ocr_text,
                'created_at': doc.created_at,
                'folder_id': doc.folder_id
            }

    st.divider()
    st.subheader(f"ğŸ“„ {doc_data['title'] or doc_data['filename']}")

    # Zusammenfassung anzeigen
    if doc_data['ai_summary']:
        st.info(f"ğŸ“ **Zusammenfassung:** {doc_data['ai_summary']}")

    # Tabs fÃ¼r verschiedene Ansichten
    tab_preview, tab_metadata, tab_edit, tab_actions = st.tabs([
        "ğŸ‘ï¸ Vorschau", "ğŸ“‹ Metadaten", "âœï¸ Bearbeiten", "âš¡ Aktionen"
    ])

    with tab_preview:
        import base64

        st.markdown("### ğŸ“„ Dokument-Vorschau")
        from utils.helpers import get_document_file_content, document_file_exists

        if doc_data['file_path'] and document_file_exists(doc_data['file_path']):
            try:
                success, result = get_document_file_content(doc_data['file_path'], doc_data.get('user_id'))
                if not success:
                    st.error(f"Fehler beim Laden: {result}")
                else:
                    # EntschlÃ¼sseln nur wenn verschlÃ¼sselt UND IV vorhanden
                    if doc_data.get('is_encrypted') and doc_data.get('encryption_iv'):
                        encryption = get_encryption_service()
                        try:
                            file_data = encryption.decrypt_file(result, doc_data['encryption_iv'], doc_data['filename'])
                        except:
                            file_data = result
                    else:
                        file_data = result

                    mime_type = doc_data['mime_type'] or ""
                    filename_lower = doc_data['filename'].lower() if doc_data['filename'] else ""

                    # PDF-Vorschau mit iframe
                    if mime_type == "application/pdf" or filename_lower.endswith(".pdf"):
                        pdf_base64 = base64.b64encode(file_data).decode('utf-8')
                        pdf_display = f'''
                        <iframe
                            src="data:application/pdf;base64,{pdf_base64}"
                            width="100%"
                            height="700px"
                            type="application/pdf"
                            style="border: 1px solid #ddd; border-radius: 5px;">
                        </iframe>
                        '''
                        st.markdown(pdf_display, unsafe_allow_html=True)

                    # Excel-Vorschau
                    elif filename_lower.endswith((".xlsx", ".xls")) or "spreadsheet" in mime_type:
                        try:
                            import pandas as pd

                            excel_file = io.BytesIO(file_data)
                            xl = pd.ExcelFile(excel_file)
                            sheet_names = xl.sheet_names

                            if len(sheet_names) > 1:
                                selected_sheet = st.selectbox(
                                    "Tabellenblatt auswÃ¤hlen",
                                    sheet_names,
                                    key=f"sheet_select_{doc_id}"
                                )
                            else:
                                selected_sheet = sheet_names[0]

                            df = pd.read_excel(excel_file, sheet_name=selected_sheet)
                            st.dataframe(df, use_container_width=True, height=500)
                            st.caption(f"ğŸ“Š {len(df)} Zeilen Ã— {len(df.columns)} Spalten")
                        except Exception as excel_err:
                            st.warning(f"Excel-Vorschau nicht mÃ¶glich: {excel_err}")

                    # Word-Vorschau (.docx)
                    elif filename_lower.endswith(".docx"):
                        try:
                            from docx import Document as DocxDocument

                            docx_file = io.BytesIO(file_data)
                            doc_content = DocxDocument(docx_file)

                            # AbsÃ¤tze extrahieren
                            full_text = []
                            for para in doc_content.paragraphs:
                                if para.text.strip():
                                    # Ãœberschriften hervorheben
                                    if para.style and para.style.name.startswith('Heading'):
                                        full_text.append(f"\n### {para.text}\n")
                                    else:
                                        full_text.append(para.text)

                            # Tabellen extrahieren
                            if doc_content.tables:
                                full_text.append("\n---\n**Tabellen:**\n")
                                for table in doc_content.tables:
                                    table_data = []
                                    for row in table.rows:
                                        row_data = [cell.text.strip() for cell in row.cells]
                                        table_data.append(" | ".join(row_data))
                                    full_text.append("\n".join(table_data))
                                    full_text.append("\n")

                            text_content = "\n".join(full_text)
                            st.markdown(text_content)
                        except ImportError:
                            st.warning("python-docx nicht installiert. Bitte installieren: pip install python-docx")
                        except Exception as word_err:
                            st.warning(f"Word-Vorschau nicht mÃ¶glich: {word_err}")

                    # Ã„ltere .doc Dateien
                    elif filename_lower.endswith(".doc"):
                        st.info("ğŸ“„ Ã„lteres Word-Format (.doc) - Bitte herunterladen und in Word Ã¶ffnen")
                        if doc_data.get('ocr_text'):
                            with st.expander("OCR-Text anzeigen"):
                                st.text_area("OCR-Text", doc_data['ocr_text'], height=300, disabled=True)

                    # Bild-Vorschau
                    elif mime_type.startswith('image/') or filename_lower.endswith((".jpg", ".jpeg", ".png", ".gif", ".webp")):
                        from PIL import Image
                        img = Image.open(io.BytesIO(file_data))
                        st.image(img, use_container_width=True)

                    # Textdateien
                    elif mime_type.startswith("text/") or filename_lower.endswith((".txt", ".csv", ".json", ".xml")):
                        try:
                            text_content = file_data.decode('utf-8')
                            if filename_lower.endswith('.csv'):
                                import pandas as pd
                                df = pd.read_csv(io.StringIO(text_content))
                                st.dataframe(df, use_container_width=True, height=500)
                            else:
                                st.code(text_content, language=None)
                        except:
                            st.warning("Textdatei konnte nicht dekodiert werden")

                    else:
                        st.info(f"ğŸ“„ Vorschau fÃ¼r {mime_type or 'unbekanntes Format'} nicht verfÃ¼gbar.")

                    # Download-Button
                    st.download_button(
                        "â¬‡ï¸ Herunterladen",
                        data=file_data,
                        file_name=doc_data['filename'],
                        mime=doc_data['mime_type'] or "application/octet-stream",
                        key="download_preview"
                    )

            except Exception as e:
                st.error(f"Fehler beim Laden: {e}")
        else:
            st.warning("Dokument-Datei nicht gefunden")

        # OCR-Text
        st.markdown("---")
        st.markdown("### ğŸ“ Erkannter Text (OCR)")
        if doc_data['ocr_text']:
            with st.expander("OCR-Text anzeigen", expanded=False):
                st.text_area("OCR-Text", doc_data['ocr_text'], height=400, disabled=True, key="ocr_preview")
        else:
            st.info("Kein OCR-Text verfÃ¼gbar")

    with tab_metadata:
        # Drei-Spalten-Layout fÃ¼r Metadaten
        col_sender, col_refs, col_finance = st.columns(3)

        with col_sender:
            st.markdown("### ğŸ“¤ Absender & Basis")
            st.write(f"**Absender:** {doc_data['sender'] or 'â€”'}")
            if doc_data['sender_address']:
                st.write(f"**Adresse:** {doc_data['sender_address']}")
            st.write(f"**Betreff:** {doc_data['subject'] or 'â€”'}")
            st.write(f"**Kategorie:** {doc_data['category'] or 'â€”'}")
            st.write(f"**Dokumentdatum:** {format_date(doc_data['document_date'])}")
            st.write(f"**Hochgeladen:** {format_date(doc_data['created_at'], True)}")

        with col_refs:
            st.markdown("### ğŸ”¢ Referenznummern")
            refs = [
                ("Aktenzeichen", doc_data['reference_number']),
                ("Kundennummer", doc_data['customer_number']),
                ("Rechnungsnr.", doc_data['invoice_number']),
                ("Vers.-Nummer", doc_data['insurance_number']),
                ("Bearbeitungsnr.", doc_data['processing_number']),
                ("Vertragsnummer", doc_data['contract_number']),
            ]
            has_refs = False
            for label, value in refs:
                if value:
                    st.write(f"**{label}:** {value}")
                    has_refs = True
            if not has_refs:
                st.caption("Keine Referenznummern erkannt")

        with col_finance:
            st.markdown("### ğŸ’° Finanzdaten")
            if doc_data.get('invoice_status'):
                from database.models import InvoiceStatus
                if doc_data['invoice_status'] == InvoiceStatus.OPEN:
                    st.error("ğŸ”´ Rechnung OFFEN")
                elif doc_data['invoice_status'] == InvoiceStatus.PAID:
                    st.success("âœ… Rechnung BEZAHLT")
                    # Zahlungsdetails anzeigen
                    if doc_data.get('invoice_paid_date'):
                        st.write(f"ğŸ“… Bezahlt am: **{format_date(doc_data['invoice_paid_date'])}**")
                    if doc_data.get('paid_with_bank_account'):
                        st.write(f"ğŸ¦ Konto: **{doc_data['paid_with_bank_account']}**")
            if doc_data['invoice_amount']:
                st.write(f"**Betrag:** {format_currency(doc_data['invoice_amount'])}")
            if doc_data['invoice_due_date']:
                st.write(f"**FÃ¤llig bis:** {format_date(doc_data['invoice_due_date'])}")
            if doc_data['iban']:
                st.code(doc_data['iban'], language=None)
                st.caption("IBAN")
            if doc_data['bic']:
                st.write(f"**BIC:** {doc_data['bic']}")
            if doc_data.get('bank_name'):
                st.write(f"**Bank:** {doc_data['bank_name']}")
            if not any([doc_data['invoice_amount'], doc_data['iban']]):
                st.caption("Keine Finanzdaten erkannt")

    with tab_edit:
        st.markdown("### âœï¸ Metadaten bearbeiten")

        with st.form("edit_metadata"):
            col_e1, col_e2 = st.columns(2)

            with col_e1:
                edit_sender = st.text_input("Absender", value=doc_data['sender'] or "")
                edit_sender_address = st.text_area("Absender-Adresse", value=doc_data['sender_address'] or "", height=100)
                edit_category = st.selectbox("Kategorie", DOCUMENT_CATEGORIES,
                    index=DOCUMENT_CATEGORIES.index(doc_data['category']) if doc_data['category'] in DOCUMENT_CATEGORIES else 0)
                edit_subject = st.text_input("Betreff", value=doc_data['subject'] or "")
                edit_doc_date = st.date_input("Dokumentdatum",
                    value=doc_data['document_date'].date() if doc_data['document_date'] else None)

            with col_e2:
                edit_ref = st.text_input("Aktenzeichen", value=doc_data['reference_number'] or "")
                edit_customer = st.text_input("Kundennummer", value=doc_data['customer_number'] or "")
                edit_invoice_nr = st.text_input("Rechnungsnummer", value=doc_data['invoice_number'] or "")
                edit_insurance = st.text_input("Versicherungsnummer", value=doc_data['insurance_number'] or "")
                edit_processing = st.text_input("Bearbeitungsnummer", value=doc_data['processing_number'] or "")
                edit_contract = st.text_input("Vertragsnummer", value=doc_data['contract_number'] or "")

            st.markdown("**Finanzdaten**")
            col_f1, col_f2, col_f3, col_f4 = st.columns(4)
            with col_f1:
                edit_amount = st.number_input("Betrag (â‚¬)", value=doc_data['invoice_amount'] or 0.0, min_value=0.0, step=0.01)
            with col_f2:
                edit_due = st.date_input("FÃ¤llig bis",
                    value=doc_data['invoice_due_date'].date() if doc_data['invoice_due_date'] else None)
            with col_f3:
                edit_iban = st.text_input("IBAN", value=doc_data['iban'] or "")
            with col_f4:
                edit_bank = st.text_input("Bank", value=doc_data.get('bank_name') or "")

            if st.form_submit_button("ğŸ’¾ Speichern", type="primary"):
                with get_db() as session:
                    doc = session.get(Document, doc_id)
                    if doc:
                        doc.sender = edit_sender or None
                        doc.sender_address = edit_sender_address or None
                        doc.category = edit_category
                        doc.subject = edit_subject or None
                        doc.title = edit_subject or doc.filename
                        doc.document_date = datetime.combine(edit_doc_date, datetime.min.time()) if edit_doc_date else None
                        doc.reference_number = edit_ref or None
                        doc.customer_number = edit_customer or None
                        doc.invoice_number = edit_invoice_nr or None
                        doc.insurance_number = edit_insurance or None
                        doc.processing_number = edit_processing or None
                        doc.contract_number = edit_contract or None
                        doc.invoice_amount = edit_amount if edit_amount > 0 else None
                        doc.invoice_due_date = datetime.combine(edit_due, datetime.min.time()) if edit_due else None
                        doc.iban = edit_iban or None
                        doc.bank_name = edit_bank or None
                        session.commit()
                        st.success("âœ… Metadaten gespeichert!")
                        st.rerun()

    with tab_actions:
        st.markdown("### âš¡ Aktionen")

        # Vorlesen-Funktion
        st.markdown("**ğŸ”Š Vorlesen**")

        from services.tts_service import get_tts_service, TTSService
        from config.settings import get_settings
        tts_settings = get_settings()

        tts_col1, tts_col2 = st.columns([2, 1])

        with tts_col1:
            tts_voice = st.selectbox(
                "Stimme",
                options=list(TTSService.VOICES.keys()),
                format_func=lambda x: TTSService.VOICES.get(x, x),
                index=list(TTSService.VOICES.keys()).index(tts_settings.tts_voice) if tts_settings.tts_voice in TTSService.VOICES else 4,
                key="tts_voice_select"
            )

        with tts_col2:
            tts_speed = st.slider("Tempo", 0.5, 2.0, tts_settings.tts_speed, 0.1, key="tts_speed_select")

        if st.button("ğŸ”Š Dokument vorlesen", use_container_width=True, type="primary"):
            if not tts_settings.openai_api_key:
                if tts_settings.tts_use_browser:
                    # Browser-TTS als Fallback
                    text_to_read = doc_data.get('ai_summary') or doc_data.get('ocr_text') or doc_data.get('subject') or "Kein Text verfÃ¼gbar"
                    tts_service = get_tts_service()
                    st.markdown(tts_service.get_browser_tts_script(text_to_read[:2000]), unsafe_allow_html=True)
                else:
                    st.warning("âš ï¸ OpenAI API nicht konfiguriert. Aktivieren Sie Browser-TTS in Einstellungen.")
            else:
                with st.spinner("Generiere Audio..."):
                    tts_service = get_tts_service()
                    result = tts_service.read_document(doc_id, tts_voice, tts_settings.tts_model, tts_speed)

                    if result.get("error"):
                        st.error(f"âŒ {result['error']}")
                    else:
                        st.audio(result["audio_bytes"], format="audio/mp3")
                        st.success("âœ… Audio generiert!")

        st.markdown("---")

        col_act1, col_act2 = st.columns(2)

        with col_act1:
            # In Aktentasche
            st.markdown("**ğŸ“‹ Aktentasche**")
            if st.button("ğŸ“‹ In Aktentasche legen", use_container_width=True):
                if 'active_cart_items' not in st.session_state:
                    st.session_state.active_cart_items = []
                if doc_id not in st.session_state.active_cart_items:
                    st.session_state.active_cart_items.append(doc_id)
                    st.success("âœ… Zur Aktentasche hinzugefÃ¼gt!")
                else:
                    st.info("Bereits in der Aktentasche")

            # Verschieben
            st.markdown("**ğŸ“‚ Ordner**")
            with get_db() as session:
                folder_tree = build_folder_tree(session, user_id)

            move_folder = st.selectbox(
                "Zielordner",
                options=[f['id'] for f in folder_tree],
                format_func=lambda x: next((f['display_name'] for f in folder_tree if f['id'] == x), ""),
                key="move_select"
            )

            # Pfad anzeigen
            selected = next((f for f in folder_tree if f['id'] == move_folder), None)
            if selected:
                st.caption(f"ğŸ“ {selected['path']}")

            col_mv, col_cp = st.columns(2)
            with col_mv:
                if st.button("ğŸ“‚ Verschieben", use_container_width=True):
                    with get_db() as session:
                        doc = session.get(Document, doc_id)
                        if doc:
                            doc.folder_id = move_folder
                            session.commit()
                            # Klassifikator lernen lassen
                            classifier = get_classifier(user_id)
                            classifier.learn_from_move(doc_id, move_folder)
                            target_name = next((f['name'] for f in folder_tree if f['id'] == move_folder), "")
                            st.success(f"âœ… Verschoben nach '{target_name}'!")
                            st.rerun()
            with col_cp:
                if st.button("ğŸ“„ Kopieren", use_container_width=True):
                    st.info("Dokument wird in Zielordner kopiert (Referenz)")

        with col_act2:
            # Teilen
            st.markdown("**ğŸ“¤ Teilen**")

            # Share-Text erstellen
            share_title = doc_data['title'] or doc_data['filename']
            share_lines = [f"ğŸ“„ {share_title}"]
            if doc_data['sender']:
                share_lines.append(f"Von: {doc_data['sender']}")
            if doc_data['category']:
                share_lines.append(f"Kategorie: {doc_data['category']}")
            if doc_data['document_date']:
                share_lines.append(f"Datum: {format_date(doc_data['document_date'])}")
            if doc_data['invoice_amount']:
                share_lines.append(f"Betrag: {format_currency(doc_data['invoice_amount'])}")
            if doc_data['iban']:
                share_lines.append(f"IBAN: {doc_data['iban']}")
            if doc_data['reference_number']:
                share_lines.append(f"Aktenzeichen: {doc_data['reference_number']}")
            share_text = "\n".join(share_lines)

            from utils.helpers import render_share_buttons
            render_share_buttons(share_title, share_text, key_prefix=f"doc_{doc_id}")

    # SchlieÃŸen-Button
    st.markdown("---")
    if st.button("âœ• SchlieÃŸen", type="secondary"):
        del st.session_state.view_document_id
        st.rerun()

# Verschieben Dialog
if 'move_document_id' in st.session_state:
    doc_id = st.session_state.move_document_id

    with st.container():
        st.divider()
        st.subheader("ğŸ“‚ Dokument verschieben")

        with get_db() as session:
            # Hierarchische Ordnerstruktur laden
            folder_tree = build_folder_tree(session, user_id)

            # Aktuellen Ordner des Dokuments ermitteln
            doc = session.get(Document, doc_id)
            current_folder_name = ""
            doc_title = ""
            if doc:
                doc_title = doc.title or doc.filename
                if doc.folder_id:
                    current_folder = session.get(Folder, doc.folder_id)
                    if current_folder:
                        current_folder_name = current_folder.name

            st.info(f"ğŸ“„ **{doc_title}** | Aktueller Ordner: **{current_folder_name or 'Kein Ordner'}**")

            # Intelligente OrdnervorschlÃ¤ge
            classifier = get_classifier(user_id)
            suggestions = classifier.suggest_folders_for_document(doc_id, limit=5)

            if suggestions:
                st.markdown("### ğŸ’¡ Empfohlene Ordner")
                for i, suggestion in enumerate(suggestions):
                    col_sug, col_btn = st.columns([4, 1])
                    with col_sug:
                        confidence_bar = "ğŸŸ¢" if suggestion['confidence'] > 0.7 else "ğŸŸ¡" if suggestion['confidence'] > 0.4 else "âšª"
                        st.markdown(f"{confidence_bar} **{suggestion['folder_name']}**")
                        st.caption(suggestion['reason'])
                    with col_btn:
                        if st.button("Hierhin", key=f"suggest_{i}_{suggestion['folder_id']}"):
                            with get_db() as move_session:
                                move_doc = move_session.get(Document, doc_id)
                                if move_doc:
                                    move_doc.folder_id = suggestion['folder_id']
                                    move_session.commit()
                                    classifier.learn_from_move(doc_id, suggestion['folder_id'])
                                    st.toast(f"âœ… Verschoben nach '{suggestion['folder_name']}'!")
                                    del st.session_state.move_document_id
                                    st.rerun()

                st.divider()

            st.markdown("### ğŸ“‚ Oder Ordner manuell wÃ¤hlen")

            # Ordnerauswahl mit Baumstruktur
            target_folder = st.selectbox(
                "Zielordner auswÃ¤hlen",
                options=[f['id'] for f in folder_tree],
                format_func=lambda x: next((f['display_name'] for f in folder_tree if f['id'] == x), ""),
                key="move_target_folder"
            )

            # Zeige vollstÃ¤ndigen Pfad
            selected_folder = next((f for f in folder_tree if f['id'] == target_folder), None)
            if selected_folder:
                st.caption(f"ğŸ“ Pfad: {selected_folder['path']}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("Verschieben", type="primary"):
                    with get_db() as move_session:
                        move_doc = move_session.get(Document, doc_id)
                        if move_doc:
                            move_doc.folder_id = target_folder
                            move_session.commit()

                            # Klassifikator lernen lassen
                            classifier.learn_from_move(doc_id, target_folder)

                            target_name = next((f['name'] for f in folder_tree if f['id'] == target_folder), "")
                            st.success(f"âœ“ Verschoben nach '{target_name}'!")
                            del st.session_state.move_document_id
                            st.rerun()
            with col2:
                if st.button("Abbrechen"):
                    del st.session_state.move_document_id
                    st.rerun()

# LÃ¶schen Dialog
if 'delete_document_id' in st.session_state:
    doc_id = st.session_state.delete_document_id

    with st.container():
        st.divider()

        # PrÃ¼fen ob Dokument bereits im Papierkorb ist
        with get_db() as session:
            doc = session.get(Document, doc_id)
            is_already_deleted = doc.is_deleted if doc else False
            doc_title = doc.title or doc.filename if doc else "Dokument"

        if is_already_deleted:
            # Im Papierkorb: EndgÃ¼ltig lÃ¶schen oder wiederherstellen
            st.warning(f"âš ï¸ '{doc_title}' ist im Papierkorb. Was mÃ¶chten Sie tun?")

            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("â™»ï¸ Wiederherstellen", type="primary"):
                    from services.trash_service import get_trash_service
                    trash_service = get_trash_service()
                    result = trash_service.restore_from_trash(doc_id, user_id)
                    if result["success"]:
                        st.success(result["message"])
                    else:
                        st.error(result["error"])
                    del st.session_state.delete_document_id
                    st.rerun()
            with col2:
                if st.button("ğŸ—‘ï¸ EndgÃ¼ltig lÃ¶schen"):
                    st.session_state.confirm_permanent_delete = doc_id
            with col3:
                if st.button("âŒ Abbrechen"):
                    del st.session_state.delete_document_id
                    st.rerun()

            # BestÃ¤tigung fÃ¼r endgÃ¼ltiges LÃ¶schen
            if st.session_state.get('confirm_permanent_delete') == doc_id:
                st.error("âš ï¸ Das Dokument wird ENDGÃœLTIG gelÃ¶scht und kann nicht wiederhergestellt werden!")
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("âœ… Ja, endgÃ¼ltig lÃ¶schen", key="confirm_perm_del"):
                        from services.trash_service import get_trash_service
                        trash_service = get_trash_service()
                        result = trash_service.permanent_delete(doc_id, user_id)
                        if result["success"]:
                            st.success(result["message"])
                        else:
                            st.error(result["error"])
                        del st.session_state.delete_document_id
                        if 'confirm_permanent_delete' in st.session_state:
                            del st.session_state.confirm_permanent_delete
                        st.rerun()
                with c2:
                    if st.button("âŒ Doch nicht", key="cancel_perm_del"):
                        del st.session_state.confirm_permanent_delete
                        st.rerun()
        else:
            # Normales LÃ¶schen: In Papierkorb verschieben
            from services.trash_service import get_trash_service
            from config.settings import get_settings
            settings = get_settings()

            st.warning(f"âš ï¸ '{doc_title}' in den Papierkorb verschieben?")
            st.info(f"ğŸ’¡ Das Dokument kann innerhalb von {settings.trash_retention_hours} Stunden wiederhergestellt werden.")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("ğŸ—‘ï¸ In Papierkorb", type="primary"):
                    trash_service = get_trash_service()
                    result = trash_service.move_to_trash(doc_id, user_id)
                    if result["success"]:
                        st.success(result["message"])
                    else:
                        st.error(result.get("error", "Fehler beim LÃ¶schen"))
                    del st.session_state.delete_document_id
                    st.rerun()
            with col2:
                if st.button("Abbrechen"):
                    del st.session_state.delete_document_id
                    st.rerun()
